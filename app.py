# -*- coding: utf-8 -*-
"""
Flask app: Phát hiện & highlight (span-level) ngôn từ xúc phạm/tiêu cực trong phản hồi sinh viên.
- POST /api/predict        -> {"text": "..."}  (phân tích 1 đoạn)
- POST /api/upload         -> multipart/form-data { file: CSV/TXT/XLSX/DOCX/PDF } (phân tích nhiều dòng)
- POST /api/export_docx    -> JSON { items: [...] } xuất DOCX có highlight
- GET  /                   -> UI

Chạy:
    pip install -r requirements.txt
    python app.py
Mở: http://127.0.0.1:5000
"""
from __future__ import annotations

import os
import re
import tempfile
from html import escape
from typing import List, Dict, Any, Optional

from flask import Flask, request, jsonify, render_template, send_file
import pandas as pd

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.linear_model import LogisticRegression

# cho export DOCX
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

# =========================
# 1) CẤU HÌNH & TỪ ĐIỂN
# =========================
profanity_list: List[str] = [
    "thối lợm",
    "ngu", "ngu ngốc", "đần", "dốt",
    "vô học", "cút đi", "cút",
    "đồ rác rưởi", "rác rưởi",
    "khốn nạn", "đáng khinh",
    "đ*m", "đ.m", "dm", "đm", "dcm",
    "vcl", "v*l", "cl",
    "đồ điên", "đồ ngu",
]

norm_dict: Dict[str, str] = {
    "dm": "địt mẹ",
    "đm": "địt mẹ",
    "dcm": "địt con mẹ",
    "cl": "cái l*n",
    "vcl": "vãi cả l*n",
    "v*l": "vãi l*n",
}

# =========================
# 2) SPAN HELPERS
# =========================
import re as _re
def _term_to_regex(term: str):
    parts = []
    for ch in term:
        if ch == '*':
            parts.append(r'[\w\.]*')
        elif ch.isspace():
            parts.append(r'\s+')
        else:
            parts.append(_re.escape(ch))
    pattern = r'(?i)\b' + ''.join(parts) + r'\b'
    return _re.compile(pattern)

LEXICON_PATTERNS = [(_w, _term_to_regex(_w)) for _w in profanity_list]

def find_spans_lexicon(original_text: str):
    spans = []
    for _, pat in LEXICON_PATTERNS:
        for m in pat.finditer(original_text):
            spans.append({"start": m.start(), "end": m.end(), "source": {"lexicon"}})
    return spans

def find_spans_abbrev(original_text: str, norm_dict: Dict[str, str]):
    spans = []
    if not norm_dict:
        return spans
    def meaning_is_profane(s: str) -> bool:
        return any(pat.search(s) for _, pat in LEXICON_PATTERNS)
    for abb, meaning in norm_dict.items():
        try:
            if meaning_is_profane(str(meaning)):
                pat_abb = _re.compile(r'(?i)\b' + _re.escape(str(abb)) + r'\b')
                for m in pat_abb.finditer(original_text):
                    spans.append({"start": m.start(), "end": m.end(), "source": {"abbrev"}})
        except Exception:
            continue
    return spans

def find_spans_ml(original_text: str, normalized_text: str, model, vectorizer, top_k: int = 3):
    spans = []
    try:
        if hasattr(model, "coef_"):
            X = vectorizer.transform([normalized_text])
            feature_names = vectorizer.get_feature_names_out()
            coefs = model.coef_[0]
            row = X.tocoo()
            contribs = []
            for idx, val in zip(row.col, row.data):
                score = val * coefs[idx]
                contribs.append((feature_names[idx], score))
            contribs.sort(key=lambda x: x[1], reverse=True)
            for tok, score in contribs[:top_k]:
                if score <= 0:
                    break
                pat_tok = _re.compile(r'(?i)\b' + _re.escape(tok) + r'\b')
                for m in pat_tok.finditer(original_text):
                    spans.append({"start": m.start(), "end": m.end(), "source": {"ml"}})
    except Exception:
        pass
    return spans

def _merge_spans(spans, text: str):
    if not spans:
        return []
    spans_sorted = sorted(spans, key=lambda s: (s["start"], s["end"]))
    merged = [spans_sorted[0]]
    for s in spans_sorted[1:]:
        last = merged[-1]
        if s["start"] <= last["end"]:
            last["end"] = max(last["end"], s["end"])
            last["source"] = set(last["source"]) | set(s["source"])
        else:
            merged.append(s)
    for m in merged:
        m["text"] = text[m["start"]:m["end"]]
        m["source"] = sorted(list(m["source"]))
    return merged

def make_highlight_html(original_text: str, spans):
    if not spans:
        return escape(original_text)
    spans_sorted = sorted(spans, key=lambda s: s["start"])
    html_parts, last = [], 0
    for s in spans_sorted:
        html_parts.append(escape(original_text[last:s["start"]]))
        tip = ", ".join(s["source"])
        html_parts.append(f"<mark title='Nguồn: {tip}'>{escape(original_text[s['start']:s['end']])}</mark>")
        last = s["end"]
    html_parts.append(escape(original_text[last:]))
    return "".join(html_parts)

# =========================
# 3) HUẤN LUYỆN / LOAD
# =========================
tfidf_vectorizer: Optional[TfidfVectorizer] = None
model: Optional[LogisticRegression] = None

def _detect_columns(df: pd.DataFrame):
    text_candidates = ["text", "content", "comment", "message", "review", "sentence"]
    label_candidates = ["label", "target", "offensive", "toxic", "is_offensive", "is_toxic", "y"]
    lower_cols = {c.lower(): c for c in df.columns}
    text_col = None
    label_col = None
    for name in text_candidates:
        if name in lower_cols:
            text_col = lower_cols[name]; break
    if text_col is None:
        obj_cols = [c for c in df.columns if df[c].dtype == 'object']
        if obj_cols:
            text_col = obj_cols[0]
    for name in label_candidates:
        if name in lower_cols:
            label_col = lower_cols[name]; break
    return text_col, label_col

def _load_user_dataset() -> Optional[pd.DataFrame]:
    candidates = [
        os.path.join(os.path.dirname(__file__), "data_train.csv"),
        "data_train.csv",
        os.path.join("/mnt/data", "data_train.csv"),
    ]
    for p in candidates:
        if os.path.exists(p):
            try:
                df = pd.read_csv(p)
                if len(df) >= 10:
                    return df
            except Exception:
                continue
    return None

def _build_synthetic_dataset() -> pd.DataFrame:
    pos = [
        "mày ngu quá", "đúng là đồ dốt", "cút đi", "đồ rác rưởi",
        "thằng này vô học", "đồ điên", "đ.m thằng kia", "vcl dở tệ",
        "thối lợm thật sự", "đáng khinh",
    ]
    neg = [
        "bạn nên cố gắng hơn", "mình góp ý nhẹ nhàng", "bài làm còn thiếu ý",
        "nên bổ sung ví dụ minh họa", "cần cải thiện cách trình bày",
        "phản hồi mang tính xây dựng", "nội dung ổn nhưng cần chi tiết hơn",
        "thầy cô sẽ hỗ trợ thêm", "cả lớp làm khá tốt", "chúc mừng bạn tiến bộ",
    ]
    df = pd.DataFrame({"text": pos + neg, "label": [1]*len(pos) + [0]*len(neg)})
    return df

def train_model():
    global tfidf_vectorizer, model
    df = _load_user_dataset()
    if df is None:
        df = _build_synthetic_dataset()
    text_col, label_col = _detect_columns(df)
    if text_col is None:
        raise RuntimeError("Không tìm thấy cột văn bản trong data_train.csv. Hãy đảm bảo có cột 'text' hoặc tương đương.")
    if label_col is None:
        def label_by_rule(s: str) -> int:
            s = str(s)
            return int(any(pat.search(s) for _, pat in LEXICON_PATTERNS))
        df["label_auto"] = df[text_col].astype(str).apply(label_by_rule)
        label_col = "label_auto"

    def normalize_text(s: str) -> str:
        t = str(s).lower().strip()
        for abb, meaning in norm_dict.items():
            pattern = r'\b' + _re.escape(abb) + r'\b'
            t = _re.sub(pattern, str(meaning), t, flags=_re.IGNORECASE)
        return t

    texts = df[text_col].astype(str).tolist()
    labels = df[label_col].astype(int).tolist()
    texts_norm = [normalize_text(x) for x in texts]

    tfidf_vectorizer = TfidfVectorizer(ngram_range=(1, 2), min_df=1, max_df=0.95)
    X = tfidf_vectorizer.fit_transform(texts_norm)

    model = LogisticRegression(max_iter=200)
    model.fit(X, labels)

train_model()

# =========================
# 4) DỰ ĐOÁN
# =========================
def preprocess_and_predict(text: str) -> Dict[str, Any]:
    original_text = str(text)
    normalized_text = original_text.lower().strip()
    if norm_dict:
        for abb, meaning in norm_dict.items():
            pattern = r'\b' + _re.escape(abb) + r'\b'
            normalized_text = _re.sub(pattern, str(meaning), normalized_text, flags=_re.IGNORECASE)

    prob_profane = None
    final_prediction = 0
    if tfidf_vectorizer is not None and model is not None:
        text_vectorized = tfidf_vectorizer.transform([normalized_text])
        raw_pred = model.predict(text_vectorized)[0]
        final_prediction = int(raw_pred)
        try:
            proba = model.predict_proba(text_vectorized)[0]
            prob_profane = round(float(proba[1]) * 100, 2)
        except Exception:
            prob_profane = None

    is_profane_by_list = any(pat.search(normalized_text) for _, pat in LEXICON_PATTERNS)

    spans = []
    spans += find_spans_lexicon(original_text)
    spans += find_spans_abbrev(original_text, norm_dict)
    spans += find_spans_ml(original_text, normalized_text, model, tfidf_vectorizer)
    spans = _merge_spans(spans, original_text)

    final_prediction = 1 if (final_prediction == 1 or is_profane_by_list or len(spans) > 0) else 0
    highlighted_html = make_highlight_html(original_text, spans)

    return {
        "normalized_text": normalized_text,
        "prediction": final_prediction,
        "probability_profane": prob_profane,
        "is_profane_by_list": bool(is_profane_by_list),
        "spans": spans,
        "highlighted_html": highlighted_html
    }

def batch_predict_texts(texts: List[str], limit: int = 200):
    items = []
    for i, t in enumerate(texts[:limit]):
        res = preprocess_and_predict(str(t))
        items.append({
            "index": i,
            "text": str(t),
            **res
        })
    labels = [f"#{it['index']+1}" for it in items]
    probs = [float(it["probability_profane"]) if isinstance(it["probability_profane"], (int, float)) else (100.0 if it["prediction"]==1 else 0.0) for it in items]
    return {
        "items": items,
        "chart": {"labels": labels, "probabilities": probs}
    }

# =========================
# 5) FLASK ROUTES
# =========================
app = Flask(__name__)

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/api/predict", methods=["POST"])
def api_predict():
    try:
        data = request.get_json(force=True) or {}
        text = data.get("text", "")
        result = preprocess_and_predict(text)
        result["chart"] = {"labels": ["Input"], "probabilities": [result["probability_profane"] if isinstance(result["probability_profane"], (int, float)) else (100.0 if result["prediction"]==1 else 0.0)]}
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 400

@app.route("/api/upload", methods=["POST"])
def api_upload():
    """
    Nhận file CSV/TXT/XLSX/DOCX/PDF và phân tích hàng loạt (tối đa 200 dòng).
    DOCX: đọc từng paragraph, PDF: trích text toàn bộ rồi tách dòng.
    """
    try:
        if "file" not in request.files:
            return jsonify({"error": "Thiếu file upload (field name: file)."}), 400
        f = request.files["file"]
        filename = f.filename or ""
        ext = filename.split(".")[-1].lower()

        df = None
        texts: List[str] = []

        if ext in {"csv"}:
            df = pd.read_csv(f)
        elif ext in {"xlsx", "xls"}:
            df = pd.read_excel(f)
        elif ext in {"txt"}:
            content = f.read().decode("utf-8", errors="ignore")
            texts = [line.strip() for line in content.splitlines() if line.strip()]
        elif ext in {"docx", "pdf"}:
            # lưu tạm để lib xử lý
            with tempfile.NamedTemporaryFile(delete=False, suffix="."+ext) as tmp:
                tmp_path = tmp.name
                f.save(tmp_path)
            try:
                if ext == "docx":
                    doc = Document(tmp_path)
                    texts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
                else:
                    from pdfminer.high_level import extract_text
                    content = extract_text(tmp_path) or ""
                    texts = [line.strip() for line in content.splitlines() if line.strip()]
            finally:
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass
        else:
            return jsonify({"error": "Định dạng không hỗ trợ. Hãy dùng CSV/TXT/XLSX/DOCX/PDF."}), 400

        if df is not None:
            text_col, _ = _detect_columns(df)
            if text_col is None:
                obj_cols = [c for c in df.columns if df[c].dtype == 'object']
                if not obj_cols:
                    return jsonify({"error": "Không tìm thấy cột văn bản trong file."}), 400
                text_col = obj_cols[0]
            texts = df[text_col].astype(str).tolist()

        result = batch_predict_texts(texts, limit=200)
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 400

@app.route("/api/export_docx", methods=["POST"])
def api_export_docx():
    """
    Nhận JSON {items:[{text:str, spans:[{start,end,text,...}], probability_profane, prediction}, ...]}
    và trả về file DOCX có highlight phần bị gắn cờ.
    """
    try:
        data = request.get_json(force=True) or {}
        items = data.get("items", [])
        if not isinstance(items, list) or not items:
            return jsonify({"error": "Không có dữ liệu items để xuất."}), 400

        doc = Document()
        doc.add_heading("BÁO CÁO PHÂN TÍCH NGÔN TỪ XÚC PHẠM", level=1)

        for it in items:
            text = str(it.get("text", ""))
            spans = it.get("spans", []) or []
            prob = it.get("probability_profane", None)
            pred = it.get("prediction", 0)

            title = f"• Dòng #{int(it.get('index', 0))+1} — Xác suất: {prob if prob is not None else 'N/A'}% — Kết luận: {'Xúc phạm' if pred==1 else 'Không'}"
            doc.add_paragraph(title)

            # paragraph có highlight
            p = doc.add_paragraph()
            last = 0
            spans_sorted = sorted(spans, key=lambda s: (s["start"], s["end"]))
            for s in spans_sorted:
                # phần trước span
                if s["start"] > last:
                    r = p.add_run(text[last:s["start"]])
                # span
                r = p.add_run(text[s["start"]:s["end"]])
                r.font.highlight_color = WD_COLOR_INDEX.YELLOW
                last = s["end"]
            # phần còn lại
            if last < len(text):
                p.add_run(text[last:])

        # lưu file tạm và trả về
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        tmp.close()
        doc.save(tmp.name)
        return send_file(tmp.name, as_attachment=True, download_name="bao_cao_highlight.docx")
    except Exception as e:
        return jsonify({"error": str(e)}), 400

@app.route("/healthz")
def healthz():
    return jsonify({"status": "ok"})

if __name__ == "__main__":
    app.run(host="127.0.0.1", port=5000, debug=True)

# =========================